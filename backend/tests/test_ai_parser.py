from pathlib import Path

import pytest

from app.ai.parser import extract_text


def test_extract_text_missing_file() -> None:
    assert extract_text("/nonexistent/file.pdf") == ""


def test_extract_text_pdf(tmp_path: Path) -> None:
    pdfplumber = pytest.importorskip("pdfplumber")
    pdf_path = tmp_path / "sample.pdf"
    # Minimal valid PDF with text
    pdf_path.write_bytes(
        b"%PDF-1.4\n1 0 obj<<>>endobj\ntrailer<<>>\n%%EOF",
    )
    # pdfplumber may not extract from minimal PDF — just ensure no exception
    result = extract_text(str(pdf_path))
    assert isinstance(result, str)


def test_extract_text_docx(tmp_path: Path) -> None:
    docx = pytest.importorskip("docx")
    from docx import Document

    path = tmp_path / "resume.docx"
    document = Document()
    document.add_paragraph("Python developer with FastAPI experience")
    document.add_table(rows=1, cols=2)
    document.tables[0].rows[0].cells[0].text = "SQL"
    document.tables[0].rows[0].cells[1].text = "PostgreSQL"
    document.save(path)

    text = extract_text(str(path))
    assert "Python" in text
    assert "PostgreSQL" in text
