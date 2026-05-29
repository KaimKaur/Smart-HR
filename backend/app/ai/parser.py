from __future__ import annotations

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_text(file_path: str) -> str:
    """Extract plain text from a PDF or DOCX resume file."""
    path = Path(file_path)
    if not path.is_file():
        logger.warning("Resume file not found: %s", file_path)
        return ""

    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            return _extract_pdf(path)
        if suffix == ".docx":
            return _extract_docx(path)
    except Exception:
        logger.exception("Failed to parse resume: %s", file_path)
        return ""
    return ""


def _extract_pdf(path: Path) -> str:
    import pdfplumber

    parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
    return "\n".join(parts).strip()


def _extract_docx(path: Path) -> str:
    from docx import Document

    document = Document(path)
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()
